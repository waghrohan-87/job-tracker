"""
Word document generator for application Q&A.
Uses python-docx (pure Python).
"""

import os
import logging
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


def _set_heading_style(paragraph, text, size=14, bold=True, colour=(44, 62, 80)):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(text)
    run.text = text
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*colour)


def _add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def generate_qa_document(
    questions_qa: list,
    job_title: str,
    company: str,
    job_url: str,
    cover_letter: str,
    output_path: str,
) -> str:
    """
    Generate a Word document containing:
    - Job details header
    - Cover letter
    - Pre-filled application Q&A

    Args:
        questions_qa: list of {question, answer} dicts from claude_engine
        job_title: str
        company: str
        job_url: str
        cover_letter: str
        output_path: full path where .docx should be saved

    Returns:
        output_path on success
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = Document()

    # ── Page margins ───────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ── Title ─────────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"Application Pack — {job_title}")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(26, 26, 26)

    # ── Job meta ──────────────────────────────────────────────────────────────
    meta = doc.add_paragraph()
    meta.add_run("Company: ").bold = True
    meta.add_run(company)
    meta.add_run("\n")
    meta.add_run("Role: ").bold = True
    meta.add_run(job_title)
    meta.add_run("\n")
    meta.add_run("Apply Link: ").bold = True
    url_run = meta.add_run(job_url)
    url_run.font.color.rgb = RGBColor(0, 102, 204)

    _add_horizontal_rule(doc)

    # ── Cover Letter ──────────────────────────────────────────────────────────
    if cover_letter:
        cl_heading = doc.add_paragraph()
        cl_run = cl_heading.add_run("COVER LETTER")
        cl_run.bold = True
        cl_run.font.size = Pt(12)
        cl_run.font.color.rgb = RGBColor(44, 62, 80)

        cl_para = doc.add_paragraph(cover_letter)
        cl_para.paragraph_format.space_after = Pt(12)

        _add_horizontal_rule(doc)

    # ── Application Questions ─────────────────────────────────────────────────
    if questions_qa:
        qa_heading = doc.add_paragraph()
        qa_run = qa_heading.add_run("APPLICATION QUESTIONS & ANSWERS")
        qa_run.bold = True
        qa_run.font.size = Pt(12)
        qa_run.font.color.rgb = RGBColor(44, 62, 80)

        doc.add_paragraph(
            "Review and edit these answers before submitting. "
            "They are pre-filled based on your resume.",
            style="Normal",
        ).runs[0].font.color.rgb = RGBColor(120, 120, 120)

        doc.add_paragraph("")

        for i, qa in enumerate(questions_qa, 1):
            # Question
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"Q{i}: {qa.get('question', '')}")
            q_run.bold = True
            q_run.font.size = Pt(10)
            q_run.font.color.rgb = RGBColor(26, 26, 26)

            # Answer
            a_para = doc.add_paragraph()
            a_run = a_para.add_run(f"A: {qa.get('answer', '')}")
            a_run.font.size = Pt(10)
            a_run.font.color.rgb = RGBColor(60, 60, 60)
            a_para.paragraph_format.left_indent = Inches(0.2)
            a_para.paragraph_format.space_after = Pt(10)

    elif not questions_qa:
        doc.add_paragraph(
            "No specific application questions were detected in the job description. "
            "The cover letter above should be sufficient.",
            style="Normal",
        )

    doc.save(output_path)
    logger.info(f"Q&A Word doc saved: {output_path}")
    return output_path
