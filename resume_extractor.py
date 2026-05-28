"""
resume_extractor.py — Run this ONCE locally after uploading a new resume to AITable.

What it does:
1. Downloads your resume attachment from AITable
2. Extracts text from it (PDF or Word)
3. Runs it through Claude to clean up formatting
4. Writes the extracted_text back to the same AITable record

Usage:
    python resume_extractor.py

Run this:
- When you first set up the system
- Any time you replace your resume with a new version
"""

import os
import sys
import logging
import tempfile
import requests
from dotenv import load_dotenv

load_dotenv()

logging.basicConfig(level=logging.INFO, format="%(asctime)s | %(levelname)s | %(message)s")
logger = logging.getLogger("resume_extractor")


def extract_text_from_pdf(path: str) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text.strip()
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return ""


def extract_text_from_docx(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())
        # Also try tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text.strip())
        return "\n".join(lines)
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""


def run():
    from utils.aitable_client import get_resume_record, update_resume_extracted_text
    from ai_engine.claude_engine import extract_resume_text_from_raw

    logger.info("Looking for active resume in AITable...")
    record = get_resume_record()

    if not record:
        logger.error("No active resume found. Make sure is_active is ticked in the resume datasheet.")
        sys.exit(1)

    record_id = record["recordId"]
    fields = record.get("fields", {})
    attachments = fields.get("resume_file", [])

    if not attachments:
        logger.error("No file attached to the active resume record. Upload your PDF or Word file first.")
        sys.exit(1)

    attachment = attachments[0]
    file_url = attachment.get("url") or attachment.get("token")
    file_name = attachment.get("name", "resume.pdf").lower()

    if not file_url:
        logger.error(f"Could not find download URL in attachment: {attachment}")
        sys.exit(1)

    logger.info(f"Downloading resume: {file_name}")

    with tempfile.TemporaryDirectory() as tmpdir:
        file_path = os.path.join(tmpdir, file_name)

        # Download
        headers = {"Authorization": f"Bearer {os.environ['AITABLE_API_TOKEN']}"}
        r = requests.get(file_url, headers=headers, timeout=30)
        r.raise_for_status()
        with open(file_path, "wb") as f:
            f.write(r.content)

        logger.info(f"Downloaded {len(r.content)} bytes")

        # Extract text
        if file_name.endswith(".pdf"):
            raw_text = extract_text_from_pdf(file_path)
        elif file_name.endswith(".docx") or file_name.endswith(".doc"):
            raw_text = extract_text_from_docx(file_path)
        else:
            logger.error(f"Unsupported file type: {file_name}. Use PDF or DOCX.")
            sys.exit(1)

        if not raw_text or len(raw_text) < 100:
            logger.error("Extracted text is too short — file may be image-based or corrupted.")
            sys.exit(1)

        logger.info(f"Extracted {len(raw_text)} characters of raw text")

        # Clean with Claude
        logger.info("Cleaning resume text with Claude...")
        clean_text = extract_resume_text_from_raw(raw_text)
        logger.info(f"Cleaned text: {len(clean_text)} characters")

        # Write back to AITable
        logger.info("Saving extracted text to AITable...")
        update_resume_extracted_text(record_id, clean_text)

        logger.info("✓ Resume text extracted and saved successfully!")
        logger.info("You can now run main.py or push to GitHub to start the automated runs.")

        # Preview
        print("\n" + "=" * 50)
        print("RESUME TEXT PREVIEW (first 500 chars):")
        print("=" * 50)
        print(clean_text[:500])
        print("...")


if __name__ == "__main__":
    run()
